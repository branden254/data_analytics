import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from matplotlib import pyplot as plt
from wordcloud import WordCloud
import seaborn as sns

# Load the data
data_path = r"C:\Users\carso\Documents\september\Level Up Tech Store\Telesales registry sample.xlsx"
df = pd.read_excel(data_path)

# Ensure date columns are parsed correctly and filter only paid sales
df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
df['Price'] = pd.to_numeric(df['Price'], errors='coerce')
df['Cost'] = pd.to_numeric(df['Cost'], errors='coerce')
df['Margin'] = pd.to_numeric(df['Margin'], errors='coerce')
df_paid = df[df['Status'] == 'PAID']

# Prepare the Word document
doc = Document()
doc.add_heading('Telesales Analysis Report - September 2024', level=1)

# 1. Profit Metrics - Table and Visual
total_revenue = df_paid['Price'].sum()
total_cost = df_paid['Cost'].sum()
total_margin = df_paid['Margin'].sum()
average_margin = df_paid['Margin'].mean()
transaction_count = len(df_paid)

profit_data = {
    'Total Revenue': total_revenue,
    'Total Cost': total_cost,
    'Total Margin': total_margin,
    'Average Margin per Sale': average_margin,
    'Number of Transactions': transaction_count
}

doc.add_heading('1. Profit Metrics', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'

for key, value in profit_data.items():
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = f"{value:,.2f} KSh"  # Format with KSh suffix

# Plot Revenue vs. Cost for Profit Metrics
plt.figure(figsize=(6, 4))
sns.barplot(x=list(profit_data.keys())[:3], y=[total_revenue, total_cost, total_margin], palette="viridis")

# Custom y-axis formatter to display values in KSh
def ksh_formatter(x, _):
    return f"{int(x):,} KSh"

plt.gca().yaxis.set_major_formatter(FuncFormatter(ksh_formatter))

plt.title("Profit Metrics")
plt.xlabel("Metric")
plt.ylabel("Amount")
plt.tight_layout()
plt.savefig("profit_metrics.png")
doc.add_picture("profit_metrics.png", width=Inches(5.5))


# 2. Daily Sales Trend
daily_sales = df_paid.groupby(df_paid['Date'].dt.date)['Price'].sum()
plt.figure(figsize=(10, 6))
daily_sales.plot(kind='line', marker='o')
plt.title('Daily Sales Trend')
plt.xlabel('Date')
plt.ylabel('Sales Amount')
plt.grid()
plt.savefig("daily_sales_trend.png")
plt.show()
doc.add_heading('3. Daily Sales Trend', level=2)
doc.add_picture('daily_sales_trend.png', width=Inches(5.5))

# 3. Top Products Analysis
top_products = df_paid.groupby('Product').agg({
    'Price': 'sum', 'Margin': 'sum'
}).sort_values('Price', ascending=False).head(10)

doc.add_heading('4. Top Products Analysis', level=2)
table = doc.add_table(rows=1 + len(top_products), cols=3)
table.style = 'Table Grid'
table.cell(0, 0).text = "Product"
table.cell(0, 1).text = "Revenue"
table.cell(0, 2).text = "Margin"

for i, (product, row) in enumerate(top_products.iterrows(), 1):
    table.cell(i, 0).text = product
    table.cell(i, 1).text = f"{row['Price']:.2f}"
    table.cell(i, 2).text = f"{row['Margin']:.2f}"



# Additional sections would continue similarly to create tables, figures, and visualizations...
# For brevity, only these sections are displayed here.
# Continuing from the previous code...


# 7. Customer Type Analysis (CUSTOMER vs RESELLER)
customer_type_analysis = df_paid.groupby('Remarks').agg({
    'Price': 'sum', 'Margin': 'sum', 'Customer Name': 'count'
}).rename(columns={'Customer Name': 'Transactions'})

doc.add_heading('7. Customer Type Analysis', level=2)
table = doc.add_table(rows=1 + len(customer_type_analysis), cols=4)
table.style = 'Table Grid'
table.cell(0, 0).text = "Customer Type"
table.cell(0, 1).text = "Sales"
table.cell(0, 2).text = "Margin"
table.cell(0, 3).text = "Transactions"

for i, (ctype, row) in enumerate(customer_type_analysis.iterrows(), 1):
    table.cell(i, 0).text = ctype
    table.cell(i, 1).text = f"{row['Price']:.2f}"
    table.cell(i, 2).text = f"{row['Margin']:.2f}"
    table.cell(i, 3).text = str(row['Transactions'])

# Plot customer type comparison
plt.figure(figsize=(8, 6))
sns.barplot(x=customer_type_analysis.index, y=customer_type_analysis['Price'], palette="viridis")
plt.title("Sales by Customer Type")
plt.xlabel("Customer Type")
plt.ylabel("Total Sales")
plt.tight_layout()
plt.savefig("customer_type_analysis.png")
doc.add_picture("customer_type_analysis.png", width=Inches(5.5))

# 7. Payment Status Analysis
payment_status = df.groupby('Status').size()
conversion_rate = df[df['Status'] == 'PAID'].shape[0] / df.shape[0] if df.shape[0] > 0 else 0

doc.add_heading('8. Payment Status Analysis', level=2)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = "Metric"
table.cell(0, 1).text = "Value"
table.cell(1, 0).text = "Total Paid Transactions"
table.cell(1, 1).text = str(df[df['Status'] == 'PAID'].shape[0])
table.cell(2, 0).text = "Conversion Rate"
table.cell(2, 1).text = f"{conversion_rate:.2%}"

# Visualizing Payment Status Distribution
plt.figure(figsize=(8, 6))
payment_status.plot(kind='pie', autopct='%1.1f%%', startangle=90, colors=['#4CAF50', '#FFC107', '#F44336'])
plt.title("Payment Status Distribution")
plt.ylabel("")  # Hide y-label for pie chart clarity
plt.tight_layout()
plt.savefig("payment_status_analysis.png")
doc.add_picture("payment_status_analysis.png", width=Inches(5.5))

# 8. High-Value Transactions Analysis (Top 10% by Value)
high_value_threshold = df_paid['Price'].quantile(0.9)
high_value_transactions = df_paid[df_paid['Price'] >= high_value_threshold]

doc.add_heading('9. High-Value Transactions Analysis (Top 10%)', level=2)
table = doc.add_table(rows=1 + len(high_value_transactions), cols=4)
table.style = 'Table Grid'
table.cell(0, 0).text = "Customer Name"
table.cell(0, 1).text = "Product"
table.cell(0, 2).text = "Price"
table.cell(0, 3).text = "Margin"

# Populate the table with high-value transaction data
for i, (_, row) in enumerate(high_value_transactions.iterrows(), 1):
    table.cell(i, 0).text = row['Customer Name']
    table.cell(i, 1).text = row['Product']
    table.cell(i, 2).text = f"{row['Price']:.2f}"
    table.cell(i, 3).text = f"{row['Margin']:.2f}"

#9. Weekly sales trend
# Ensure 'Date' column is of datetime type
df_paid['Date'] = pd.to_datetime(df_paid['Date'])

# Step 2: Define the weeks in September 2024
week_starts = pd.date_range(start='2024-09-01', end='2024-09-30', freq='W-SUN')  # Weeks start on Sundays
week_labels = {i+1: f'Week {i+1}' for i in range(len(week_starts))}

# Step 3: Create a new column for the corresponding week
def get_week(date):
    for i, week_start in enumerate(week_starts):
        if week_start <= date < (week_start + pd.Timedelta(days=7)):
            return i + 1  # Return week number (1-5)
    return None

df_paid['Custom Week'] = df_paid['Date'].apply(get_week)

# Step 4: Group by custom week and aggregate performance metrics
weekly_performance = df_paid.groupby('Custom Week').agg({
    'Price': 'sum', 
    'Cost': 'sum', 
    'Margin': 'sum', 
    'Customer Name': 'count'
}).rename(columns={'Customer Name': 'Transactions'}).reindex(week_labels.keys())  # Reorder weeks

# Fill NaN values with zeros for weeks without transactions
weekly_performance.fillna(0, inplace=True)

# Step 5: Add a heading for the analysis in the Word document
doc.add_heading('10. Weekly Performance Analysis', level=2)
table = doc.add_table(rows=1 + len(weekly_performance), cols=5)
table.style = 'Table Grid'
table.cell(0, 0).text = "Week"
table.cell(0, 1).text = "Sales"
table.cell(0, 2).text = "Cost"
table.cell(0, 3).text = "Margin"
table.cell(0, 4).text = "Transactions"

# Populate the table with the weekly performance data
for i, (week, row) in enumerate(weekly_performance.iterrows(), 1):
    table.cell(i, 0).text = week_labels[week]  # Use custom week labels
    table.cell(i, 1).text = f"{row['Price']:.2f}"
    table.cell(i, 2).text = f"{row['Cost']:.2f}"
    table.cell(i, 3).text = f"{row['Margin']:.2f}"
    table.cell(i, 4).text = str(int(row['Transactions']))  # Ensure transactions are displayed as integers

# Step 6: Plot for Weekly Sales Trend
plt.figure(figsize=(10, 6))
weekly_performance['Price'].plot(kind='line', marker='o', color='purple')
plt.title("Weekly Sales Trend for September 2024")
plt.xlabel("Week")
plt.ylabel("Sales Amount")

# Adding an empty label for x=0 and aligning the labels correctly
plt.xticks(ticks=range(len(week_labels) + 1), labels=[''] + list(week_labels.values()), rotation=45)  # Adding an empty label for 0
plt.grid()
plt.tight_layout()
plt.savefig("weekly_sales_trend.png")
doc.add_picture("weekly_sales_trend.png", width=Inches(5.5))

# 10. Word cloud of most frequently mentioned products or brands
text = ' '.join(df['Product'].dropna())
wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
plt.figure(figsize=(10, 5))
plt.imshow(wordcloud, interpolation='bilinear')
plt.axis('off')
plt.title('Most Frequently Mentioned Products')
plt.savefig('product_wordcloud.png')
doc.add_picture("product_wordcloud.png", width=Inches(5.5))
plt.close()

# Save the Document
doc.save("Telesales_Analysis_Report_September2.docx")
