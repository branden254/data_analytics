import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from docx import Document
from docx.shared import Inches
from io import BytesIO
import numpy as np
from matplotlib.patches import Circle, RegularPolygon
from matplotlib.path import Path
from matplotlib.projections.polar import PolarAxes
from matplotlib.projections import register_projection
from matplotlib.spines import Spine
from matplotlib.transforms import Affine2D

def load_data(file_path):
    return pd.read_excel(file_path)

def create_document():
    doc = Document()
    doc.add_heading('Level Up Competitor Analysis Summary', 0)
    return doc

def add_overall_price_competitiveness(doc, df):
    doc.add_heading('1. Overall Price Competitiveness', level=1)
    df['Price Difference'] = df['Price'] - df['Least Price']
    df['Price Competitiveness'] = df['Price Difference'].apply(
        lambda x: 'Higher' if x > 0 else ('Lower' if x < 0 else 'Equal'))
    
    summary = df['Price Competitiveness'].value_counts().reset_index()
    summary.columns = ['Price Comparison', 'Count']
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Price Comparison'
    hdr_cells[1].text = 'Count'
    for _, row in summary.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Price Comparison'])
        row_cells[1].text = str(row['Count'])
    
    doc.add_paragraph('Overview of how Flex prices compare with competitors across all categories.')

def add_price_distribution_by_category(doc, df):
    doc.add_heading('2. Price Distribution by Category', level=1)
    plt.figure(figsize=(12, 6))
    sns.boxplot(x='Category', y='Price', data=df)
    plt.title('Price Distribution by Category')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This plot shows the price distribution across different product categories.')

def add_top_competitors(doc, df):
    doc.add_heading('3. Top Competitors Offering Lowest Prices', level=1)
    top_competitors = df['Company with Least Price'].value_counts().nlargest(10)
    plt.figure(figsize=(10, 6))
    sns.barplot(y=top_competitors.index, x=top_competitors.values)
    plt.title('Top 10 Competitors Offering Lowest Prices')
    plt.tight_layout()
    plt.show()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This chart shows the top competitors who most frequently offer the lowest prices across all categories.')

def add_price_comparison_scatter(doc, df):
    doc.add_heading('4. Flex vs Average Competitor Prices', level=1)
    plt.figure(figsize=(10, 6))
    plt.scatter(df['AVERAGE COMPETITOR PRICE'], df['Price'], alpha=0.5)
    plt.title('Flex vs Average Competitor Prices')
    plt.xlabel('Average Competitor Price')
    plt.ylabel('Flex Price')
    plt.plot([df['AVERAGE COMPETITOR PRICE'].min(), df['AVERAGE COMPETITOR PRICE'].max()], 
             [df['AVERAGE COMPETITOR PRICE'].min(), df['AVERAGE COMPETITOR PRICE'].max()], 'r--')
    plt.tight_layout()
    plt.show()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This scatter plot compares Flex prices to average competitor prices across all products.')

def add_profit_margin_analysis(doc, df):
    doc.add_heading('5. Profit Margin Analysis', level=1)
    df['Estimated Profit Margin %'] = ((df['Price'] - df['Least Price']) / df['Price'] * 100).clip(lower=0)
    profit_margin = df.groupby('Category')['Estimated Profit Margin %'].agg(['mean', 'min', 'max']).reset_index()
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Average Profit Margin %'
    hdr_cells[2].text = 'Min Profit Margin %'
    hdr_cells[3].text = 'Max Profit Margin %'
    
    for _, row in profit_margin.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Category'])
        row_cells[1].text = f"{row['mean']:.2f}%"
        row_cells[2].text = f"{row['min']:.2f}%"
        row_cells[3].text = f"{row['max']:.2f}%"
    
    doc.add_paragraph('This table summarizes the estimated profit margins for each category.')
def add_price_trend_analysis(doc, df):
    doc.add_heading('6. Price Trend Analysis', level=1)
    df['Price Difference %'] = (df['Price'] - df['AVERAGE COMPETITOR PRICE']) / df['AVERAGE COMPETITOR PRICE'] * 100
    trend = df.groupby('Category')['Price Difference %'].mean().sort_values(ascending=False)
    
    # Create a horizontal lollipop chart
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.hlines(y=trend.index, xmin=0, xmax=trend.values, color='skyblue')
    ax.plot(trend.values, trend.index, "o")
    
    ax.set_xlabel('Average Price Difference %')
    ax.set_title('Average Price Difference % by Category', fontsize=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.yaxis.set_ticks_position('none')
    
    # Add percentage labels
    for i, v in enumerate(trend.values):
        ax.text(v, i, f' {v:.1f}%', va='center')
    
    plt.tight_layout()
    plt.show()
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This lollipop chart shows the average percentage difference between Flex prices and competitor prices for each category. Positive values indicate Flex prices are higher, negative values indicate they are lower.')

def add_market_share_analysis(doc, df):
    doc.add_heading('7. Category Distribution Analysis', level=1)
    category_distribution = df['Category'].value_counts()
    
    # Show top 10 categories and group the rest as "Others"
    top_categories = category_distribution.nlargest(10)
    others = pd.Series({'Others': category_distribution.iloc[10:].sum()})
    category_distribution = pd.concat([top_categories, others])
    
    # Create a donut chart
    fig, ax = plt.subplots(figsize=(10, 10))
    colors = plt.cm.Spectral(np.linspace(0, 1, len(category_distribution)))
    
    wedges, texts, autotexts = ax.pie(category_distribution.values, 
                                      labels=category_distribution.index, 
                                      autopct='%1.1f%%',
                                      pctdistance=0.85,
                                      colors=colors)
    
    # Create the donut center
    centre_circle = plt.Circle((0, 0), 0.70, fc='white')
    fig.gca().add_artist(centre_circle)
    
    # Adjust text positions to avoid overlapping
    for i, autotext in enumerate(autotexts):
        ang = (wedges[i].theta2 + wedges[i].theta1) / 2.
        y = np.sin(np.deg2rad(ang))
        x = np.cos(np.deg2rad(ang))
        horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
        connectionstyle = f"angle,angleA=0,angleB={ang}"
        autotext.set_position((1.2*x, 1.2*y))
        autotext.set_ha(horizontalalignment)
        texts[i].set_position((1.3*x, 1.3*y))
        texts[i].set_ha(horizontalalignment)
        
    ax.set_title('Distribution of Products by Category', fontsize=15)
    plt.tight_layout()
    plt.show()
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This donut chart illustrates the distribution of products across different categories. Each slice represents a category, with the size indicating the proportion of products in that category. The top 10 categories are shown individually, with the rest grouped as "Others".')

def add_price_range_analysis(doc, df):
    doc.add_heading('8. Price Range Analysis', level=1)
    df['Price Range'] = pd.cut(df['Price'], bins=[0, 100, 500, 1000, 5000, float('inf')],
                               labels=['0-100', '101-500', '501-1000', '1001-5000', '5000+'])
    price_range_dist = df['Price Range'].value_counts().sort_index()
    
    # Create a violin plot
    fig, ax = plt.subplots(figsize=(12, 6))
    sns.violinplot(x='Price Range', y='Price', data=df, ax=ax)
    
    ax.set_ylabel('Price (KSH)')
    ax.set_title('Distribution of Prices Across Price Ranges', fontsize=15)
    ax.set_yscale('log')  # Use log scale for better visibility
    
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This violin plot shows the distribution of prices across different price ranges. The width of each "violin" represents the frequency of products at that price point, while the height shows the price range.')

def add_competitor_price_correlation(doc, df):
    doc.add_heading('9. Competitor Price Correlation', level=1)
    correlation = df.groupby('Category').apply(lambda x: x['Price'].corr(x['AVERAGE COMPETITOR PRICE'])).sort_values(ascending=False)
    
    # Create a scatter plot with regression lines
    fig, ax = plt.subplots(figsize=(12, 8))
    
    categories = correlation.nlargest(5).index  # Top 5 categories by correlation
    for category in categories:
        category_data = df[df['Category'] == category]
        sns.regplot(x='AVERAGE COMPETITOR PRICE', y='Price', data=category_data, 
                    scatter=True, fit_reg=True, label=category, scatter_kws={'alpha':0.3})
    
    ax.set_xlabel('Average Competitor Price (KSH)')
    ax.set_ylabel('Flex Price (KSH)')
    ax.set_title('Flex Price vs Average Competitor Price\nTop 5 Categories by Correlation', fontsize=15)
    ax.legend()
    
    plt.tight_layout()
    plt.show()
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph('This scatter plot shows the relationship between Flex prices and average competitor prices for the top 5 categories with the highest price correlation. Each line represents the trend for a specific category.')

def main():
    file_path = r"C:\Users\carso\Documents\september\Flex\PRICE LISTS\Flex Cometitor analysis.xlsx"
    df = load_data(file_path)
    doc = create_document()
    
    add_overall_price_competitiveness(doc, df)
    add_price_distribution_by_category(doc, df)
    add_top_competitors(doc, df)
    add_price_comparison_scatter(doc, df)
    add_profit_margin_analysis(doc, df)
    add_price_trend_analysis(doc, df)
    add_market_share_analysis(doc, df)
    add_price_range_analysis(doc, df)
    add_competitor_price_correlation(doc, df)
    
    doc.save('Flex_Generalized_Competitor_Analysis4.docx')
    print("Generalized competitor analysis summary generated: Flex_Generalized_Competitor_Analysis4.docx")

if __name__ == "__main__":
    main()
