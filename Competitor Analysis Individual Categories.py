import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the data
file_path = r"C:\Users\Administrator\Documents\WORK\August\week 2\competitor analysis laptops.xlsx"
laptops_df = pd.read_excel(file_path, sheet_name='Laptops')

# Create a Word document
doc = Document()
doc.add_heading('Laptop Competitor Analysis', 0)

# 1. Price Comparison Analysis
doc.add_heading('1. Price Comparison Analysis', level=1)
laptops_df['Price Difference vs Lowest'] = laptops_df['GENSPACE PRICES'] - laptops_df['Least Price']
price_comparison = laptops_df[['BRAND', 'PRODUCT', 'GENSPACE PRICES', 'Least Price', 'Price Difference vs Lowest']].sort_values(by='Price Difference vs Lowest', ascending=False).head(10)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(price_comparison.columns):
    hdr_cells[i].text = column_name
for row in price_comparison.itertuples():
    row_cells = table.add_row().cells
    for i in range(5):
        row_cells[i].text = str(row[i+1])
doc.add_paragraph('This analysis compares GenSpace prices to the lowest competitor prices, highlighting significant differences.')

# 2. Competitor Pricing Distribution by Brand
doc.add_heading('2. Competitor Pricing Distribution by Brand', level=1)
plt.figure(figsize=(10, 6))
sns.boxplot(x='BRAND', y='Least Price', data=laptops_df)
plt.title('Competitor Pricing Distribution by Brand')
plt.xticks(rotation=45)
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This visualization shows the distribution of competitor prices for each brand.')

# 3. Price Recommendations
doc.add_heading('3. Price Recommendations', level=1)
laptops_df['% Difference vs Lowest'] = (laptops_df['GENSPACE PRICES'] - laptops_df['Least Price']) / laptops_df['Least Price'] * 100
laptops_df['Recommendation'] = laptops_df['% Difference vs Lowest'].apply(lambda x: 'Decrease Price' if x > 20 else ('Increase Price' if x < -20 else 'Maintain Price'))
recommendations = laptops_df[['BRAND', 'PRODUCT', '% Difference vs Lowest', 'Recommendation']].head(10)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(recommendations.columns):
    hdr_cells[i].text = column_name
for row in recommendations.itertuples():
    row_cells = table.add_row().cells
    for i in range(4):
        row_cells[i].text = str(row[i+1])
doc.add_paragraph('This analysis provides recommendations to increase, decrease, or maintain prices based on the percentage difference from the lowest competitor price.')

# 4. Comparison to Average Competitor Price
doc.add_heading('4. Comparison to Average Competitor Price', level=1)
laptops_df['% Difference vs Average'] = (laptops_df['GENSPACE PRICES'] - laptops_df['Average Competitor price ']) / laptops_df['Average Competitor price '] * 100
avg_comparison = laptops_df[['BRAND', 'PRODUCT', 'GENSPACE PRICES', 'Average Competitor price ', '% Difference vs Average']].head(10)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(avg_comparison.columns):
    hdr_cells[i].text = column_name
for row in avg_comparison.itertuples():
    row_cells = table.add_row().cells
    for i in range(5):
        row_cells[i].text = str(row[i+1])
doc.add_paragraph('This analysis compares GenSpace prices to the average competitor price for each product.')

# 5. Best Price Offerings Analysis
doc.add_heading('5. Best Price Offerings Analysis', level=1)
best_price_offerings = laptops_df['Company with Least Price'].value_counts().head(10)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Company'
hdr_cells[1].text = 'Lowest Price Offers'
for company, count in best_price_offerings.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(company)
    row_cells[1].text = str(count)
doc.add_paragraph('This analysis identifies companies that most frequently offer the lowest prices.')

# 6. Brand Price Analysis
doc.add_heading('6. Brand Price Analysis', level=1)
brand_price_analysis = laptops_df.groupby('BRAND').agg({
    'GENSPACE PRICES': 'mean',
    'Least Price': 'mean',
    'Average Competitor price ': 'mean'
}).reset_index()
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(brand_price_analysis.columns):
    hdr_cells[i].text = column_name
for row in brand_price_analysis.itertuples():
    row_cells = table.add_row().cells
    for i in range(4):
        row_cells[i].text = str(row[i+1])
doc.add_paragraph('This analysis provides an overview of how different brands position themselves in terms of pricing.')

# 7. Profit Margin Analysis
doc.add_heading('7. Profit Margin Analysis', level=1)
laptops_df['Profit Margin'] = (laptops_df['GENSPACE PRICES'] - laptops_df['Least Price']) / laptops_df['GENSPACE PRICES'] * 100
profit_margin = laptops_df[['BRAND', 'PRODUCT', 'Profit Margin']].head(10)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(profit_margin.columns):
    hdr_cells[i].text = column_name
for row in profit_margin.itertuples():
    row_cells = table.add_row().cells
    for i in range(3):
        row_cells[i].text = str(row[i+1])
doc.add_paragraph('This analysis helps in understanding the profitability of each product and identifying areas for potential price adjustments.')

# 8. Price Evaluation Summary
doc.add_heading('8. Price Evaluation Summary', level=1)
price_evaluation_summary = laptops_df['Price Evaluation'].value_counts()
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Price Evaluation'
hdr_cells[1].text = 'Count'
for evaluation, count in price_evaluation_summary.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(evaluation)
    row_cells[1].text = str(count)
doc.add_paragraph('This gives an overall picture of how GenSpace\'s pricing strategy is performing.')

# 9. Price Trend Analysis (commented out as in original code)
doc.add_heading('9. Price Trend Analysis', level=1)
doc.add_paragraph('This analysis would show how prices have changed over time for both GenSpace and competitors. It requires date data which is not available in the current dataset.')

# 10. Competitor Performance Analysis
doc.add_heading('10. Competitor Performance Analysis', level=1)
competitor_performance = laptops_df.groupby('Company with Least Price').size().reset_index(name='Lowest Price Offers')
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(competitor_performance.columns):
    hdr_cells[i].text = column_name
for row in competitor_performance.itertuples():
    row_cells = table.add_row().cells
    for i in range(2):
        row_cells[i].text = str(row[i+1])
doc.add_paragraph('This identifies which competitors are most frequently undercutting GenSpace on price.')

# Save the document
doc.save('laptop_competitor_analysis.docx')
print("Word document generated: laptop_competitor_analysis_2024.docx")
