import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from datetime import timedelta

# Load the dataset
file_path = r"C:\Users\carso\Documents\october\WEEK 3\mondays october.xlsx"
data = pd.read_excel(file_path)

# Convert the 'Date' column to datetime
data['Date'] = pd.to_datetime(data['Date'], errors='coerce')

# Define the Mondays and their corresponding weeks
mondays = pd.to_datetime(['2024-08-05', '2024-08-12','2024-08-19','2024-08-26',])

# Create a Word document
doc = Document()
doc.add_heading('Telesales Analysis: Monday vs Specific Days of the Week', 0)

# Define function to create week ranges based on Monday
def get_week_range(monday):
    start = monday
    end = monday + timedelta(days=6)
    return (start, end)

# Function to get the day of the week from a date
def get_day_of_week(date):
    return date.strftime('%A')

# Iterate through each Monday and compare it to other days in the week
for monday in mondays:
    start, end = get_week_range(monday)
    weekly_data = data[(data['Date'] >= start) & (data['Date'] <= end)]
    
    # Split data into days of the week
    monday_data = weekly_data[weekly_data['Date'] == monday]
    day_data = {get_day_of_week(day): weekly_data[weekly_data['Date'] == day] 
                for day in pd.date_range(start=start, end=end)}

    # Initialize analysis variables
    inquiries = {}
    sales = {}
    outcomes = {}

    # Collect data for each day of the week
    for day_name, day_df in day_data.items():
        # Analysis 1: Inquiries (based on non-null 'Remarks')
        inquiries[day_name] = day_df['Remarks'].notna().sum()

        # Analysis 2: Sales (defined as Price * Qty Ordered)
        sales[day_name] = (day_df['Price'] * day_df['Qty Ordered']).sum()

        # Analysis 3: Call Outcomes (distribution of outcomes)
        outcomes[day_name] = day_df['Call Outcome'].value_counts()

    # Plot 1: Inquiries comparison (Monday vs other days of the week)
    plt.figure()
    plt.bar(inquiries.keys(), inquiries.values(), color=['blue' if day == 'Monday' else 'orange' for day in inquiries.keys()])
    plt.title(f'Inquiries: {monday.date()} Week ({start.date()} to {end.date()})')
    plt.ylabel('Number of Inquiries')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f'inquiries_week_{monday.date()}.png')
    plt.close()

    # Plot 2: Sales comparison (Monday vs other days of the week)
    plt.figure()
    plt.bar(sales.keys(), sales.values(), color=['green' if day == 'Monday' else 'red' for day in sales.keys()])
    plt.title(f'Sales: {monday.date()} Week ({start.date()} to {end.date()})')
    plt.ylabel('Total Sales (Price * Qty Ordered)')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f'sales_week_{monday.date()}.png')
    plt.close()

    # Plot 3: Call outcomes comparison (Monday vs other days of the week)
    plt.figure()
    valid_days = [day for day, outcome in outcomes.items() if not outcome.empty]  # Ensure we only plot days with data
    if valid_days:
        monday_colors = {
            'Follow-up needed': 'green',    # Monday-specific colors
            'Out Of Stock': 'green',
            'Pending': 'green',
            'Closed Sale': 'green',
            'Unlisted': 'green'
        }
        other_day_color = 'gray'

        for day in valid_days:
            outcome = outcomes[day]
            
            if day == 'Monday':
                outcome_colors = [monday_colors.get(outcome_type, 'green') for outcome_type in outcome.index]
            else:
                outcome_colors = [other_day_color for _ in outcome.index]
            
            outcome.plot(kind='bar', alpha=0.7, color=outcome_colors, label=day, width=0.8 / len(valid_days), position=valid_days.index(day))

        plt.title(f'Call Outcomes: {monday.date()} Week ({start.date()} to {end.date()})')
        plt.ylabel('Number of Calls')
        plt.legend()
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig(f'call_outcomes_week_{monday.date()}.png')
        plt.close()
    else:
        print(f"No call outcome data available for the week starting {monday.date()}")

    # Add content to Word document for the current week
    doc.add_heading(f'Week of {monday.date()}: {start.date()} to {end.date()}', level=1)
    
    # Inquiries Section
    doc.add_heading('Inquiries', level=2)
    doc.add_paragraph(f'Inquiries breakdown for the week:')
    for day, count in inquiries.items():
        doc.add_paragraph(f'{day}: {count} inquiries')
    doc.add_picture(f'inquiries_week_{monday.date()}.png', width=Inches(5))
    
    # Sales Section
    doc.add_heading('Sales', level=2)
    doc.add_paragraph(f'Sales breakdown for the week (calculated as Price * Qty Ordered):')
    for day, total_sales in sales.items():
        doc.add_paragraph(f'{day}: {total_sales:.2f} total sales')
    doc.add_picture(f'sales_week_{monday.date()}.png', width=Inches(5))

    # Call Outcomes Section
    doc.add_heading('Call Outcomes', level=2)
    doc.add_paragraph('Call outcomes breakdown for each day:')
    for day, outcome in outcomes.items():
        doc.add_paragraph(f'{day}:')
        for outcome_type, count in outcome.items():
            doc.add_paragraph(f'    {outcome_type}: {count}')
    doc.add_picture(f'call_outcomes_week_{monday.date()}.png', width=Inches(5))

# Save the document after all weeks have been processed
doc.save('Telesales_Analysis_Monday_vs_Days.docx')

print('Analysis complete! The Word document has been generated with all weeks.')
