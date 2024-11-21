import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the data
file_path = r"C:\Users\carso\Documents\september\Flex\PRICE LISTS\Flex Cometitor analysis.xlsx"
laptops_df = pd.read_excel(file_path, sheet_name='Sheet1')

# Create a Word document
doc = Document()
doc.add_heading('Flex Toners and Other Accessories Competitor Analysis', 0)
# 1. Price Competitiveness Overview
doc.add_heading('1. Price Competitiveness Overview', level=1)
laptops_df['Price Difference vs Lowest'] = laptops_df['Price'] - laptops_df['Least Price']

# Add a column to categorize if GenSpace's price is higher, lower, or equal to the lowest competitor price
laptops_df['Price Competitiveness'] = laptops_df.apply(lambda row: 
    'Higher' if row['Price'] > row['Least Price'] else
    'Lower' if row['Price'] < row['Least Price'] else 'Equal', axis=1)

# Count the number of products in each category
price_competitive_summary = laptops_df['Price Competitiveness'].value_counts().reset_index()
price_competitive_summary.columns = ['Price Comparison', 'Count']

# Add this data to the document as a table
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Price Comparison'
hdr_cells[1].text = 'Count'

for row in price_competitive_summary.itertuples(index=False):
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = str(row[1])

doc.add_paragraph('This analysis categorizes products based on whether Flex prices are higher, lower, or equal to the lowest competitor prices, providing an overview of pricing competitiveness.')

# 2. Average Prices by Brand
doc.add_heading('2. Average Prices by Category', level=1)
plt.figure(figsize=(10, 6))
sns.barplot(x='Category', y='Price', data=laptops_df)
plt.title('Average Flex\'s by Category')
plt.xticks(rotation=45)
plt.ylabel('Average Price')
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png', bbox_inches='tight')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This bar plot shows the average Flex prices for each brand, allowing for easy comparison of pricing strategies across different manufacturers.')

# 3. Price Distribution
doc.add_heading('3. Price Distribution', level=1)
plt.figure(figsize=(10, 6))
sns.histplot(data=laptops_df, x='Price', kde=True)
plt.title('Distribution of Flex Prices')
plt.xlabel('Price')
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png', bbox_inches='tight')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This histogram displays the distribution of Flex prices, showing the most common price ranges and any outliers.')

# 4. GenSpace vs Average Competitor Prices
doc.add_heading('4. Flex  vs Average Competitor Prices', level=1)
plt.figure(figsize=(10, 6))
plt.scatter(laptops_df['Average Competitor price '], laptops_df['Price'])
plt.title('Flex  Prices vs Average Competitor Prices')
plt.xlabel('Average Competitor Price')
plt.ylabel('Flex  Price')
plt.plot([laptops_df['Average Competitor price '].min(), laptops_df['Average Competitor price '].max()], 
         [laptops_df['Average Competitor price '].min(), laptops_df['Average Competitor price '].max()], 
         'r--')
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png', bbox_inches='tight')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This scatter plot compares Level Up  prices to average competitor prices. Points above the red line indicate where Flex prices are higher than the average competitor price.')

# 5. Price Evaluation Summary
doc.add_heading('5. Price Evaluation Summary', level=1)
plt.figure(figsize=(8, 8))
price_eval_summary = laptops_df['Price Evaluation'].value_counts()
plt.pie(price_eval_summary, labels=price_eval_summary.index, autopct='%1.1f%%')
plt.title('Price Evaluation Summary')
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png', bbox_inches='tight')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This pie chart summarizes the price evaluation categories, showing the proportion of products that are priced competitively, high, or low.')

# 6. Top Competitors Offering Lowest Prices
doc.add_heading('6. Top Competitors Offering Lowest Prices', level=1)
plt.figure(figsize=(10, 6))
top_competitors = laptops_df['Company with Least Price'].value_counts().nlargest(10)
sns.barplot(y=top_competitors.index, x=top_competitors.values)
plt.title('Top 10 Competitors Offering Lowest Prices')
plt.xlabel('Number of Lowest Price Offers')
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png', bbox_inches='tight')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This horizontal bar chart identifies the top competitors who most frequently offer the lowest prices, helping to pinpoint the main price competitors.')

# 7. Correlation Heatmap
doc.add_heading('7. Correlation Heatmap of Price Variables', level=1)
plt.figure(figsize=(10, 8))
corr_matrix = laptops_df[['Price', 'Least Price', 'Average Competitor price ', 'Price Difference vs Lowest']].corr()
sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
plt.title('Correlation Heatmap of Price Variables')
img_buffer = BytesIO()
plt.savefig(img_buffer, format='png', bbox_inches='tight')
img_buffer.seek(0)
doc.add_picture(img_buffer, width=Inches(6))
doc.add_paragraph('This heatmap shows the correlation between different price variables, helping to understand the relationships between Flex prices, competitor prices, and price differences.')

# 8. Brand Market Positioning
doc.add_heading('8. Category Market Positioning', level=1)
brand_positioning = laptops_df.groupby('Category').agg({
    'Price': ['mean', 'min', 'max'],
    'Least Price': 'mean',
    'Average Competitor price ': 'mean'
}).reset_index()
brand_positioning.columns = ['Category', 'Avg GenSpace Price', 'Min GenSpace Price', 'Max GenSpace Price', 'Avg Least Price', 'Avg Competitor Price']
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
for i, column_name in enumerate(brand_positioning.columns):
    table.cell(0, i).text = column_name
for row in brand_positioning.itertuples(index=False):
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = f"{value:.2f}" if isinstance(value, float) else str(value)
doc.add_paragraph('This table shows how different categories position themselves in terms of price range and target market.')
# 9. Pricing Strategy Effectiveness
doc.add_heading('9. Pricing Strategy Effectiveness', level=1)

# Handle missing values and calculate price difference
laptops_df['Price Difference %'] = laptops_df.apply(lambda row: 
    0 if pd.isnull(row['Average Competitor price ']) or row['Average Competitor price '] == 0
    else (row['Price'] - row['Average Competitor price ']) / row['Average Competitor price '] * 100, 
    axis=1)

# Group by brand and calculate statistics
price_strategy = laptops_df.groupby('Category').agg({
    'Price Difference %': ['mean', 'count'],
    'Price': 'count'  # Total product count
}).reset_index()

price_strategy.columns = ['Category', 'Avg Price Difference %', 'Comparable Product Count', 'Total Product Count']

# Calculate the percentage of products with competitors
price_strategy['% Products with Competitors'] = (price_strategy['Comparable Product Count'] / price_strategy['Total Product Count'] * 100).round(2)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, column_name in enumerate(price_strategy.columns):
    table.cell(0, i).text = column_name

for row in price_strategy.itertuples(index=False):
    cells = table.add_row().cells
    cells[0].text = str(row[0])  # BRAND
    cells[1].text = f"{row[1]:.2f}%"  # Avg Price Difference %
    cells[2].text = str(row[2])  # Comparable Product Count
    cells[3].text = str(row[3])  # Total Product Count
    cells[4].text = f"{row[4]}%"  # % Products with Competitors

doc.add_paragraph('This analysis shows how Flex\'s pricing strategy compares to competitors across different brands. '
                  'It includes the average price difference percentage, the number of products with competitor prices, '
                  'the total number of products, and the percentage of products that have competitors.')
# 10. Profit Margin Analysis
doc.add_heading('10. Profit Margin Analysis', level=1)
laptops_df['Estimated Profit Margin %'] = (laptops_df['Price'] - laptops_df['Least Price']) / laptops_df['Price'] * 100
profit_margin = laptops_df.groupby('Category')['Estimated Profit Margin %'].agg(['mean', 'min', 'max']).reset_index()
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for i, column_name in enumerate(profit_margin.columns):
    table.cell(0, i).text = column_name
for row in profit_margin.itertuples(index=False):
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = f"{value:.2f}%" if i > 0 else str(value)
doc.add_paragraph('This analysis estimates profit margins by Category, helping to identify areas for potential price adjustments.')

# Save the document
doc.save('Flex_competitor_analysis_updated_1.docx')
print("Updated Word document generated: Flex_competitor_analysis_updated.docx")
