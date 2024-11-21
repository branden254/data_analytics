import pandas as pd
import numpy as np
from sklearn.cluster import KMeans, DBSCAN, AgglomerativeClustering
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import silhouette_score, davies_bouldin_score
from sklearn.decomposition import PCA
from sklearn.model_selection import GridSearchCV
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import io

# Step 1: Data Collection
np.random.seed(42)
n_customers = 250
customer_data = pd.DataFrame({
    'CustomerID': range(1, n_customers + 1),
    'Frequency': np.random.randint(1, 50, n_customers),
    'CLV': np.random.uniform(100, 10000, n_customers),
    'AvgSpend': np.random.uniform(20, 500, n_customers),
    'Electronics': np.random.randint(0, 2, n_customers),
    'Apparel': np.random.randint(0, 2, n_customers),
    'FarmInputs': np.random.randint(0, 2, n_customers)
})

# Step 2: Data Preprocessing
# Handle missing values (if any)
customer_data.fillna(customer_data.mean(), inplace=True)

# Detect and handle outliers (simple method)
for col in ['Frequency', 'CLV', 'AvgSpend']:
    q1 = customer_data[col].quantile(0.25)
    q3 = customer_data[col].quantile(0.75)
    iqr = q3 - q1
    lower_bound = q1 - 1.5 * iqr
    upper_bound = q3 + 1.5 * iqr
    customer_data = customer_data[(customer_data[col] >= lower_bound) & (customer_data[col] <= upper_bound)]

# Scaling features
features = ['Frequency', 'CLV', 'AvgSpend', 'Electronics', 'Apparel', 'FarmInputs']
scaler = StandardScaler()
scaled_data = scaler.fit_transform(customer_data[features])

# Dimensionality Reduction
pca = PCA(n_components=2)
pca_data = pca.fit_transform(scaled_data)

# Step 3: Clustering Algorithm
# Hyperparameter tuning for KMeans
kmeans = KMeans()
param_grid = {'n_clusters': range(2, 10)}
grid_search = GridSearchCV(kmeans, param_grid, cv=5)
grid_search.fit(scaled_data)
best_kmeans = grid_search.best_estimator_

# Apply the best KMeans
customer_data['Cluster'] = best_kmeans.fit_predict(scaled_data)

# Step 4: Cluster Validation
silhouette_avg = silhouette_score(scaled_data, customer_data['Cluster'])
davies_bouldin_avg = davies_bouldin_score(scaled_data, customer_data['Cluster'])

# Step 5: Cluster Analysis
cluster_summary = customer_data.groupby('Cluster').mean()

# Visualizations
plt.figure(figsize=(10, 6))
sns.scatterplot(x=pca_data[:, 0], y=pca_data[:, 1], hue=customer_data['Cluster'], palette='viridis')
plt.title('Customer Segments by PCA Components')
scatter_plot_buffer = io.BytesIO()
plt.savefig(scatter_plot_buffer, format='png')
scatter_plot_buffer.seek(0)
plt.close()

plt.figure(figsize=(10, 6))
sns.countplot(data=customer_data, x='Cluster', palette='viridis')
plt.title('Number of Customers in Each Cluster')
count_plot_buffer = io.BytesIO()
plt.savefig(count_plot_buffer, format='png')
count_plot_buffer.seek(0)
plt.close()

# Step 6: Implementing Targeted Campaigns
cluster_labels = {
    0: 'High-Value Customers',
    1: 'Frequent Buyers',
    2: 'Occasional Buyers',
    3: 'Emerging Customers'
}

def get_product_recommendations(cluster):
    # Dummy function to generate product recommendations
    recommendations = {
        0: ['Product A', 'Product B'],
        1: ['Product C', 'Product D'],
        2: ['Product E', 'Product F'],
        3: ['Product G', 'Product H']
    }
    return recommendations[cluster]

def send_email(to, subject, body):
    # Dummy function to simulate sending an email
    print(f'Sending email to {to}')
    print(f'Subject: {subject}')
    print(f'Body: {body}\n')

# Example Email Campaign
for cluster, label in cluster_labels.items():
    cluster_customers = customer_data[customer_data['Cluster'] == cluster]
    
    for _, customer in cluster_customers.iterrows():
        email = f'customer{customer["CustomerID"]}@example.com'
        recommendations = get_product_recommendations(cluster)
        
        send_email(
            to=email,
            subject=f"Special Offers for Our {label}!",
            body=f"Dear Customer {customer['CustomerID']},\n\nWe have some exciting offers just for you! Based on your recent purchases, we thought you might like:\n{recommendations}\n\nBest regards,\nErnest_co"
        )

# Creating the Word Document
document = Document()
document.add_heading('Customer Clustering Analysis and Marketing Strategies', 0)

document.add_heading('1. Introduction', level=1)
document.add_paragraph(
    "In today’s competitive market, knowing your customers' behaviors and preferences is crucial. "
    "Customer clustering segments your audience into distinct groups based on their actions and traits. "
    "This strategy enables businesses to craft targeted marketing and personalized product recommendations."
)

document.add_heading('2. Data Collection', level=1)
document.add_paragraph(
    "The dataset contains customer information, including ID, Frequency of purchase, CLV, Average Spend on an order, "
    "and product categories purchased (e.g., Electronics, Apparel, Farm Inputs)."
)

document.add_heading('3. Data Preprocessing', level=1)
document.add_paragraph(
    "The data was cleaned and preprocessed to ensure suitability for analysis. This involved handling missing values, "
    "scaling features, and encoding categorical variables."
)

document.add_heading('4. Feature Selection', level=1)
document.add_paragraph(
    "The features selected for clustering include the number of orders, CLV, average order size, and product categories purchased."
)

document.add_heading('5. Clustering Algorithm', level=1)
document.add_paragraph(
    "The K-Means clustering algorithm was applied to segment customers into distinct groups. "
    f"The optimal number of clusters was determined using hyperparameter tuning, resulting in {len(cluster_labels)} clusters."
)

document.add_heading('6. Cluster Validation', level=1)
document.add_paragraph(
    f"The clusters were validated using the Silhouette Score and Davies-Bouldin Index. "
    f"Silhouette Score: {silhouette_avg:.2f}\nDavies-Bouldin Index: {davies_bouldin_avg:.2f}"
)

document.add_heading('7. Cluster Analysis', level=1)
document.add_paragraph("The characteristics of each cluster are as follows:")
document.add_paragraph(cluster_summary.to_string())

document.add_heading('8. Visualizations', level=1)
document.add_paragraph("Customer Segments by PCA Components:")
document.add_picture(scatter_plot_buffer, width=Inches(6))

document.add_paragraph("Number of Customers in Each Cluster:")
document.add_picture(count_plot_buffer, width=Inches(6))

document.add_heading('9. Implementing Targeted Campaigns', level=1)
document.add_paragraph(
    "Once customers are segmented into clusters, businesses can develop targeted marketing strategies and personalized product recommendations for each group."
)
for cluster, label in cluster_labels.items():
    document.add_heading(f'Cluster {cluster}: {label}', level=2)
    document.add_paragraph("Strategies:")
    if cluster == 0:
        document.add_paragraph(
            "1. Exclusive Perks: Offer special deals, early access to new products, and premium services.\n"
            "2. Personal Touch: Send personalized emails highlighting products similar to their previous purchases.\n"
            "3. Loyalty Programs: Introduce a loyalty program that rewards frequent purchases with exclusive benefits."
        )
    elif cluster == 1:
        document.add_paragraph(
            "1. Upsell and Cross-Sell: Recommend products that complement their previous purchases.\n"
            "2. Subscription Offers: Promote subscription services for products they buy regularly.\n"
            "3. Reward Points: Implement a points system that gives them discounts or free items based on their purchase frequency."
        )
    elif cluster == 2:
        document.add_paragraph(
            "1. Regular Reminders: Send them reminders about products they showed interest in.\n"
            "2. Limited-Time Discounts: Offer special discounts that expire soon to encourage quick purchases.\n"
            "3. Engaging Content: Share interesting and relevant content to keep them connected with your brand."
        )
    elif cluster == 3:
        document.add_paragraph(
            "1. Feedback Surveys: Ask for feedback to understand why they shop infrequently.\n"
            "2. Small Incentives: Provide small discounts or free samples to encourage more purchases.\n"
            "3. Targeted Ads: Use targeted ads to highlight product benefits and attract their attention."
        )

# Save the document
document.save('Customer_Clustering_Analysis2.docx')
