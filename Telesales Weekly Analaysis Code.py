import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Load data
df = pd.read_excel(r"C:\Users\Administrator\Documents\WORK\september\week 2\Telecalls iwst week september.xlsx")

# Convert Date column to datetime and filter for the last week
df['Date'] = pd.to_datetime(df['Date'])
df['Date'] = df['Date'].dt.strftime('%Y-%m-%d')

# Define the date range for the last week (from Monday 01/07/2024 to Saturday 06/07/2024)
date_range = pd.date_range(start='2024-9-02', end='2024-9-07')
date_mapping = {date.strftime('%Y-%m-%d'): day for date, day in zip(date_range, ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'])}
last_week = df[df['Date'].isin(date_mapping.keys())]
last_week['Day of Week'] = last_week['Date'].map(date_mapping)
last_week['Sales'] = last_week['Price'] * last_week['Qty Ordered']

# Create a Word document
doc = Document()
doc.add_heading('Telesales Registry Analysis (Last Week)', 0)

# 1. Weekly Sales Trend Analysis
doc.add_heading('1. Weekly Sales Trend Analysis', level=1)
doc.add_paragraph('This graph shows the total sales for each day of the last week. Total sales are calculated by multiplying the price of each product by the quantity ordered.')
daily_sales = last_week.groupby('Day of Week')['Sales'].sum().reindex(['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'])
plt.figure(figsize=(10, 6))
daily_sales.plot(kind='line', marker='o')
plt.title('Daily Sales Trend (Last Week)')
plt.xlabel('Day of Week')
plt.ylabel('Sales (KSH)')
plt.grid(True)
plt.savefig('daily_sales_trend.png')
plt.show()
doc.add_picture('daily_sales_trend.png', width=Inches(5))
plt.close()

# List daily sales
doc.add_paragraph('Daily Sales:')
for day, sales in daily_sales.items():
    doc.add_paragraph(f'{day}: KSH {sales:.2f}')

# 2. Call Outcome Distribution (Pie Chart)
doc.add_heading('2. Call Outcome Distribution', level=1)
doc.add_paragraph('This pie chart shows the distribution of call outcomes over the last week. It provides an overview of the results of the telesales calls.')
call_outcome_counts = last_week['Call Outcome'].value_counts()
# Plotting the Call Outcome Distribution as a pie chart
plt.figure(figsize=(10, 6))
call_outcome_counts.plot(kind='pie', autopct='%1.1f%%', colors=plt.cm.Paired.colors)
plt.title('Call Outcome Distribution (Last Week)')
plt.ylabel('')  # Hide the y-label for a cleaner look
plt.savefig('call_outcome_distribution_pie.png')
plt.show()
doc.add_picture('call_outcome_distribution_pie.png', width=Inches(5))
plt.close()

# Listing call outcome counts
doc.add_paragraph('Call Outcome Counts:')
for outcome, count in call_outcome_counts.items():
    doc.add_paragraph(f'{outcome}: {count}')

# 3. Top Selling Products
doc.add_heading('3. Top Selling Products', level=1)
doc.add_paragraph('This bar chart shows the top 5 products by total sales over the last week. Total sales are calculated by multiplying the price of each product by the quantity ordered.')
product_sales = last_week.groupby('Product')['Sales'].sum().nlargest(5)
plt.figure(figsize=(10, 6))
product_sales.plot(kind='bar')
plt.title('Top 5 Products by Sales (Last Week)')
plt.xlabel('Product')
plt.ylabel('Sales (KSH)')
plt.xticks(rotation=45)
plt.grid(True)
plt.savefig('top_selling_products.png')
plt.show()
doc.add_picture('top_selling_products.png', width=Inches(5))
plt.close()

# List top 5 selling products with sales
doc.add_paragraph('Top 5 Selling Products with Sales:')
for product, sales in product_sales.items():
    doc.add_paragraph(f'{product}: KSH {sales:.2f}')

# List the top 5 selling products with quantities
top_5_products = last_week.groupby('Product')['Qty Ordered'].sum().nlargest(5)
doc.add_paragraph('Top 5 Selling Products with Quantities:')
for product, qty in top_5_products.items():
    doc.add_paragraph(f'{product}: {qty}')

# 4. Sales by Media
doc.add_heading('4. Sales by Media', level=1)
doc.add_paragraph('This pie chart shows the distribution of sales across different media channels over the last week. Sales are calculated by multiplying the price of each product by the quantity ordered.')
media_sales = last_week.groupby('Media')['Sales'].sum()
plt.figure(figsize=(10, 6))
media_sales.plot(kind='pie', autopct='%1.1f%%')
plt.title('Sales by Media (Last Week)')
plt.ylabel('')
plt.savefig('sales_by_media.png')
plt.show()
doc.add_picture('sales_by_media.png', width=Inches(5))
plt.close()

# List media sales
doc.add_paragraph('Sales by Media:')
for media, sales in media_sales.items():
    doc.add_paragraph(f'{media}: KSH {sales:.2f}')

# 5. Average Order Value (AOV)
doc.add_heading('5. Average Order Value (AOV)', level=1)
doc.add_paragraph('The Average Order Value (AOV) is calculated by dividing the total sales by the total number of orders.')
total_sales = last_week['Sales'].sum()
total_orders = last_week['Qty Ordered'].sum()
aov = total_sales / total_orders
doc.add_paragraph(f'Average Order Value (Last Week): KSH {aov:.2f}')

# 6. Top 5 Repetitive Numbers
doc.add_heading('6. Top 5 Repetitive Numbers', level=1)
doc.add_paragraph('This analysis identifies the top 5 most frequently called numbers over the last week.')
top_5_numbers = last_week['No'].value_counts().nlargest(5)
doc.add_paragraph('Top 5 Repetitive Numbers:')
for number, count in top_5_numbers.items():
    doc.add_paragraph(f'{number}: {count} calls')

# 7. Profit Margin Analysis
doc.add_heading('7. Profit Margin Analysis', level=1)
doc.add_paragraph('The total profit margin is calculated as the sum of the margins for all sales made over the last week.')
total_margin = last_week['Margin'].sum()
doc.add_paragraph(f'Total Profit Margin (Last Week): KSH {total_margin:.2f}')

# 8. Follow-up Needed Analysis
doc.add_heading('8. Follow-up Needed Analysis', level=1)
doc.add_paragraph('This analysis identifies the number of calls that need follow-up over the last week.')
total_calls = len(last_week)
follow_up_needed = last_week[last_week['Call Outcome'] == 'Follow Up-Needed'].shape[0]
follow_up_percentage = (follow_up_needed / total_calls) * 100
doc.add_paragraph(f'Number of Follow-up Needed Calls (Last Week): {follow_up_needed} out of {total_calls} ({follow_up_percentage:.2f}%)')

# 9. Out of Stock Analysis
doc.add_heading('9. Out of Stock Analysis', level=1)
doc.add_paragraph('This analysis identifies the products that were out of stock over the last week.')

# Total number of products (including duplicates) in the entire dataset for last week
total_products = last_week['Product'].count()

# Filter out-of-stock products
out_of_stock_products = last_week[last_week['Call Outcome'] == 'Out Of Stock']['Product']

# Count the number of out-of-stock products
out_of_stock_count = out_of_stock_products.count()

# Calculate the percentage of out-of-stock products
out_of_stock_percentage = (out_of_stock_count / total_products) * 100

# Add the analysis to the Word document
doc.add_paragraph(f'Number of Out of Stock Products (Last Week): {out_of_stock_count} out of {total_products} ({out_of_stock_percentage:.2f}%)')
doc.add_paragraph('Out of Stock Products (Last Week):')
for product in out_of_stock_products.unique():
    doc.add_paragraph(product)

# 10. Sales by Day of the Week
doc.add_heading('10. Sales by Day of the Week', level=1)
doc.add_paragraph('This bar chart shows the sales distribution by day of the week over the last week. Sales are calculated by multiplying the price of each product by the quantity ordered.')
sales_by_day = last_week.groupby('Day of Week')['Sales'].sum().reindex(['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'])
plt.figure(figsize=(10, 6))
sales_by_day.plot(kind='bar')
plt.title('Sales by Day of the Week (Last Week)')
plt.xlabel('Day of the Week')
plt.ylabel('Sales (KSH)')
plt.xticks(rotation=45)
plt.grid(True)
plt.savefig('sales_by_day_of_week.png')
plt.show()
doc.add_picture('sales_by_day_of_week.png', width=Inches(5))
plt.close()

# List sales by day of the week
doc.add_paragraph('Sales by Day of the Week:')
for day, sales in sales_by_day.items():
    doc.add_paragraph(f'{day}: KSH {sales:.2f}')

# 11. Day with the Most Calls
doc.add_heading('11. Day with the Most Calls', level=1)
doc.add_paragraph('This analysis identifies the day with the most calls over the last week.')
day_most_calls = last_week['Day of Week'].value_counts().idxmax()
doc.add_paragraph(f'Day with the Most Calls: {day_most_calls}')

# 12. Day with the Least Calls
doc.add_heading('12. Day with the Least Calls', level=1)
doc.add_paragraph('This analysis identifies the day with the least calls over the last week.')
day_least_calls = last_week['Day of Week'].value_counts().idxmin()
doc.add_paragraph(f'Day with the Least Calls: {day_least_calls}')

# 13. Day with the Most Profit
doc.add_heading('13. Day with the Most Profit', level=1)
doc.add_paragraph('This analysis identifies the day with the most profit over the last week. Profit is calculated as the difference between sales and cost.')
profit_by_day = last_week.groupby('Day of Week')['Margin'].sum().reindex(['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'])
day_most_profit = profit_by_day.idxmax()
doc.add_paragraph(f'Day with the Most Profit: {day_most_profit}')

# 14. Most Frequent Inquired Products
doc.add_heading('14. Most Frequent Inquired Products', level=1)
doc.add_paragraph('This analysis identifies the top 10 most frequently mentioned products in the product section over the last week.')
most_frequent_products = last_week['Product'].value_counts().nlargest(10)
doc.add_paragraph('Top 10 Most Frequent Products:')
for product, count in most_frequent_products.items():
    doc.add_paragraph(f'{product}: {count} times')
# 15. Category Sales Analysis
doc.add_heading('15. Category Sales Analysis', level=1)
doc.add_paragraph('This analysis examines the total sales by category, providing insights into which product categories generated the most revenue during the week.')

# Grouping data by Category and summing up Sales
category_sales = last_week.groupby('Category')['Sales'].sum()

# Plotting the Category Sales
plt.figure(figsize=(10, 6))
category_sales.plot(kind='bar', color='skyblue')
plt.title('Total Sales by Category (Last Week)')
plt.xlabel('Category')
plt.ylabel('Sales (KSH)')
plt.xticks(rotation=45)
plt.grid(True)
plt.savefig('category_sales_analysis.png')
plt.show()
doc.add_picture('category_sales_analysis.png', width=Inches(5))
plt.close()

# Listing the sales by category
doc.add_paragraph('Sales by Category:')
for category, sales in category_sales.items():
    doc.add_paragraph(f'{category}: KSH {sales:.2f}')
# 16. Inquiries per Category
doc.add_heading('16. Inquiries per Category', level=1)
doc.add_paragraph('This analysis shows the number of inquiries for each product category over the last week.')

# Count inquiries per category
inquiries_per_category = last_week['Category'].value_counts()

# Plotting inquiries per category
plt.figure(figsize=(10, 6))
inquiries_per_category.plot(kind='bar', color='lightgreen')
plt.title('Inquiries per Category (Last Week)')
plt.xlabel('Category')
plt.ylabel('Number of Inquiries')
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
plt.grid(axis='y')
plt.savefig('inquiries_per_category.png')
plt.show()
plt.close()

doc.add_picture('inquiries_per_category.png', width=Inches(5))

# List inquiries per category
doc.add_paragraph('Inquiries per Category:')
for category, count in inquiries_per_category.items():
    doc.add_paragraph(f'{category}: {count} inquiries')
# 17. Inquiries vs Number of Sales per Category
doc.add_heading('17. Inquiries vs Number of Sales per Category', level=1)
doc.add_paragraph('This analysis compares the number of inquiries to the number of successful sales for each product category over the last week.')

# Prepare data
category_inquiries = last_week['Category'].value_counts()
category_sales_count = last_week[last_week['Sales'] > 0]['Category'].value_counts()

# Combine inquiries and sales count data
category_comparison = pd.DataFrame({'Inquiries': category_inquiries, 'Sales': category_sales_count}).fillna(0)
category_comparison = category_comparison.sort_values('Inquiries', ascending=False)

# Plotting inquiries vs number of sales per category
fig, ax1 = plt.subplots(figsize=(12, 6))

# Plot inquiries as bars
ax1.bar(category_comparison.index, category_comparison['Inquiries'], alpha=0.7, color='lightblue', label='Inquiries')
ax1.set_xlabel('Category')
ax1.set_ylabel('Number of Inquiries', color='blue')
ax1.tick_params(axis='y', labelcolor='blue')
ax1.set_xticklabels(category_comparison.index, rotation=45, ha='right')

# Create a secondary y-axis for sales
ax2 = ax1.twinx()
ax2.plot(category_comparison.index, category_comparison['Sales'], color='red', marker='o', label='Sales')
ax2.set_ylabel('Number of Sales', color='red')
ax2.tick_params(axis='y', labelcolor='red')

plt.title('Inquiries vs Number of Sales per Category (Last Week)')
fig.legend(loc='upper right', bbox_to_anchor=(1,1), bbox_transform=ax1.transAxes)

plt.tight_layout()
plt.savefig('inquiries_vs_sales_count_per_category.png')
plt.show()
plt.close()

doc.add_picture('inquiries_vs_sales_count_per_category.png', width=Inches(5))

# List inquiries vs number of sales per category
doc.add_paragraph('Inquiries vs Number of Sales per Category:')
for category in category_comparison.index:
    inquiries = int(category_comparison.loc[category, 'Inquiries'])
    sales = int(category_comparison.loc[category, 'Sales'])
    doc.add_paragraph(f'{category}: {inquiries} inquiries, {sales} sales')

# Calculate and display conversion rate
doc.add_paragraph('Conversion Rates:')
for category in category_comparison.index:
    inquiries = category_comparison.loc[category, 'Inquiries']
    sales = category_comparison.loc[category, 'Sales']
    if inquiries > 0:
        conversion_rate = (sales / inquiries) * 100
        doc.add_paragraph(f'{category}: {conversion_rate:.2f}%')
    else:
        doc.add_paragraph(f'{category}: N/A (no inquiries)')

# Save the Word document
doc.save('september_Telesales_Registry_Analysis_Week_1.docx')
